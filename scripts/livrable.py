#!/usr/bin/env python3
"""
livrable.py — Générateur de documents AKOMA DIGITAL LTD
Lit un fichier markdown avec frontmatter YAML et produit un .docx à l'identité AKOMA.

Usage:
    python3 livrable.py fichier.md
    python3 livrable.py fichier.md --sortie /chemin/sortie/
    python3 livrable.py fichier.md --sortie /chemin/sortie/ --pdf
"""

import argparse
import os
import re
import subprocess
import sys
from pathlib import Path

try:
    import yaml
except ImportError:
    print("ERREUR: Le module 'pyyaml' est requis. Installez-le avec: pip install pyyaml")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERREUR: Le module 'python-docx' est requis. Installez-le avec: pip install python-docx")
    sys.exit(1)

# ============================================================
# CHARTE GRAPHIQUE AKOMA
# ============================================================
COULEUR_TERRACOTTA = RGBColor(0xB5, 0x64, 0x2F)   # #B5642F
COULEUR_ESPRESSO   = RGBColor(0x5E, 0x3A, 0x1E)   # #5E3A1E — titres
COULEUR_CREME      = RGBColor(0xF5, 0xF0, 0xE8)   # #F5F0E8 — lignes alternées
COULEUR_LIN        = RGBColor(0xE8, 0xE0, 0xD5)   # #E8E0D5 — lignes alternées paires
COULEUR_BLANC      = RGBColor(0xFF, 0xFF, 0xFF)
COULEUR_TEXTE      = RGBColor(0x1A, 0x1A, 0x1A)
POLICE_PRINCIPALE  = "Arial"
AKOMA_EMAIL        = "akoma.digital.ltd@gmail.com"
AKOMA_NOM          = "AKOMA DIGITAL LTD"
AKOMA_SLOGAN       = "Structurer. Transformer. Piloter l'impact."


def rgb_to_hex(rgb: RGBColor) -> str:
    return f"{rgb.red:02X}{rgb.green:02X}{rgb.blue:02X}"


def set_cell_background(cell, hex_color: str):
    """Définit la couleur de fond d'une cellule de tableau."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def ajouter_entete(doc: Document, titre_doc: str):
    """Ajoute l'en-tête AKOMA."""
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False

    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Ligne terracotta en haut
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), rgb_to_hex(COULEUR_TERRACOTTA))
    pBdr.append(bottom)
    pPr.append(pBdr)

    run = p.add_run(f"{AKOMA_NOM} · {AKOMA_EMAIL}")
    run.font.name = POLICE_PRINCIPALE
    run.font.size = Pt(8)
    run.font.color.rgb = COULEUR_ESPRESSO
    run.font.bold = True

    if titre_doc:
        run2 = p.add_run(f"    |    {titre_doc}")
        run2.font.name = POLICE_PRINCIPALE
        run2.font.size = Pt(8)
        run2.font.color.rgb = COULEUR_TERRACOTTA


def ajouter_pied_de_page(doc: Document):
    """Ajoute le pied de page AKOMA."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Ligne terracotta en haut du pied de page
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "6")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), rgb_to_hex(COULEUR_TERRACOTTA))
    pBdr.append(top)
    pPr.append(pBdr)

    run = p.add_run(f"{AKOMA_NOM} · {AKOMA_EMAIL} · {AKOMA_SLOGAN}")
    run.font.name = POLICE_PRINCIPALE
    run.font.size = Pt(7)
    run.font.color.rgb = COULEUR_ESPRESSO

    # Numéro de page
    p2 = footer.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_pg = p2.add_run()
    run_pg.font.size = Pt(7)
    run_pg.font.color.rgb = COULEUR_TERRACOTTA
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run_pg._r.append(fldChar1)
    run_pg._r.append(instrText)
    run_pg._r.append(fldChar2)


def configurer_styles(doc: Document):
    """Configure les styles de base du document."""
    # Style Normal
    style_normal = doc.styles["Normal"]
    style_normal.font.name = POLICE_PRINCIPALE
    style_normal.font.size = Pt(10.5)
    style_normal.font.color.rgb = COULEUR_TEXTE

    # Style Heading 1
    h1 = doc.styles["Heading 1"]
    h1.font.name = POLICE_PRINCIPALE
    h1.font.size = Pt(16)
    h1.font.color.rgb = COULEUR_ESPRESSO
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)

    # Style Heading 2
    h2 = doc.styles["Heading 2"]
    h2.font.name = POLICE_PRINCIPALE
    h2.font.size = Pt(13)
    h2.font.color.rgb = COULEUR_TERRACOTTA
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(6)

    # Style Heading 3
    h3 = doc.styles["Heading 3"]
    h3.font.name = POLICE_PRINCIPALE
    h3.font.size = Pt(11)
    h3.font.color.rgb = COULEUR_ESPRESSO
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)


def ajouter_separateur(doc: Document):
    """Ajoute un filet terracotta."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), rgb_to_hex(COULEUR_TERRACOTTA))
    pBdr.append(bottom)
    pPr.append(pBdr)


def rendre_tableau_markdown(doc: Document, lignes: list[str]):
    """Convertit un tableau markdown en tableau Word AKOMA."""
    # Extraire les lignes de données (ignorer la ligne de séparation ---)
    data_lines = []
    for ligne in lignes:
        if re.match(r"^\|[-| ]+\|$", ligne.strip()):
            continue  # ligne de séparation
        cellules = [c.strip() for c in ligne.strip().strip("|").split("|")]
        data_lines.append(cellules)

    if not data_lines:
        return

    nb_cols = max(len(row) for row in data_lines)
    # Normaliser toutes les lignes au même nombre de colonnes
    for row in data_lines:
        while len(row) < nb_cols:
            row.append("")

    table = doc.add_table(rows=len(data_lines), cols=nb_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, row_data in enumerate(data_lines):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.name = POLICE_PRINCIPALE
            run.font.size = Pt(9.5)

            if i == 0:
                # En-tête : fond terracotta, texte blanc gras
                run.font.bold = True
                run.font.color.rgb = COULEUR_BLANC
                set_cell_background(cell, rgb_to_hex(COULEUR_TERRACOTTA))
            elif i % 2 == 1:
                # Lignes impaires : crème
                run.font.color.rgb = COULEUR_TEXTE
                set_cell_background(cell, rgb_to_hex(COULEUR_CREME))
            else:
                # Lignes paires : lin
                run.font.color.rgb = COULEUR_TEXTE
                set_cell_background(cell, rgb_to_hex(COULEUR_LIN))

    doc.add_paragraph()  # espacement après le tableau


def parser_frontmatter(contenu: str) -> tuple[dict, str]:
    """Extrait le frontmatter YAML du contenu markdown."""
    match = re.match(r"^---\s*\n(.*?)\n---\s*\n", contenu, re.DOTALL)
    if match:
        try:
            meta = yaml.safe_load(match.group(1)) or {}
        except yaml.YAMLError:
            meta = {}
        corps = contenu[match.end():]
    else:
        meta = {}
        corps = contenu
    return meta, corps


def traiter_formatage_inline(run_parent, texte: str, doc: Document) -> None:
    """Applique le formatage gras/italique inline dans un paragraphe."""
    # Pattern: **gras**, *italique*, `code`
    pattern = re.compile(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)")
    parties = pattern.split(texte)
    for partie in parties:
        if partie.startswith("**") and partie.endswith("**"):
            run = run_parent.add_run(partie[2:-2])
            run.bold = True
        elif partie.startswith("*") and partie.endswith("*"):
            run = run_parent.add_run(partie[1:-1])
            run.italic = True
        elif partie.startswith("`") and partie.endswith("`"):
            run = run_parent.add_run(partie[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        else:
            run_parent.add_run(partie)


def convertir_markdown_en_docx(contenu_md: str, doc: Document, meta: dict):
    """Convertit le corps markdown en éléments Word."""
    lignes = contenu_md.split("\n")
    i = 0
    titre_doc = meta.get("titre", "")

    # Ajouter le titre principal du document
    if titre_doc:
        p = doc.add_heading(titre_doc, level=1)
        p.runs[0].font.color.rgb = COULEUR_ESPRESSO

    # Bloc méta sous le titre
    reference = meta.get("reference", "")
    date_doc = meta.get("date", "")
    client = meta.get("client", "")
    if any([reference, date_doc, client]):
        meta_parts = []
        if client:
            meta_parts.append(f"Client : {client}")
        if date_doc:
            meta_parts.append(f"Date : {date_doc}")
        if reference:
            meta_parts.append(f"Réf. : {reference}")
        p = doc.add_paragraph("  ·  ".join(meta_parts))
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = COULEUR_TERRACOTTA
        p.runs[0].font.italic = True
        doc.add_paragraph()

    ajouter_separateur(doc)

    tableau_buffer = []

    while i < len(lignes):
        ligne = lignes[i]

        # Ignorer la ligne d'en-tête markdown (# titre) si identique au frontmatter titre
        if ligne.startswith("# ") and ligne[2:].strip() == titre_doc:
            i += 1
            continue

        # Détecter début de tableau markdown
        if ligne.strip().startswith("|"):
            tableau_buffer.append(ligne)
            i += 1
            while i < len(lignes) and lignes[i].strip().startswith("|"):
                tableau_buffer.append(lignes[i])
                i += 1
            rendre_tableau_markdown(doc, tableau_buffer)
            tableau_buffer = []
            continue

        # Titres
        if ligne.startswith("### "):
            p = doc.add_heading(ligne[4:].strip(), level=3)
        elif ligne.startswith("## "):
            texte_h2 = ligne[3:].strip()
            p = doc.add_heading(texte_h2, level=2)
            # Filet terracotta sous H2
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), rgb_to_hex(COULEUR_TERRACOTTA))
            pBdr.append(bottom)
            pPr.append(pBdr)
        elif ligne.startswith("# "):
            p = doc.add_heading(ligne[2:].strip(), level=1)

        # Citation (blockquote)
        elif ligne.startswith("> "):
            texte = ligne[2:].strip()
            p = doc.add_paragraph(style="Quote")
            p.clear()
            run = p.add_run(texte)
            run.font.name = POLICE_PRINCIPALE
            run.font.size = Pt(10.5)
            run.font.italic = True
            run.font.color.rgb = COULEUR_ESPRESSO
            pPr = p._p.get_or_add_pPr()
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), "720")
            pPr.append(ind)
            pBdr = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "12")
            left.set(qn("w:space"), "4")
            left.set(qn("w:color"), rgb_to_hex(COULEUR_TERRACOTTA))
            pBdr.append(left)
            pPr.append(pBdr)

        # Listes à puces
        elif ligne.strip().startswith("- ") or ligne.strip().startswith("* "):
            indent = len(ligne) - len(ligne.lstrip())
            texte = ligne.strip()[2:]
            p = doc.add_paragraph(style="List Bullet")
            p.clear()
            traiter_formatage_inline(p, texte, doc)
            for run in p.runs:
                run.font.name = POLICE_PRINCIPALE
                run.font.size = Pt(10.5)

        # Listes numérotées
        elif re.match(r"^\s*\d+\.\s", ligne):
            texte = re.sub(r"^\s*\d+\.\s", "", ligne)
            p = doc.add_paragraph(style="List Number")
            p.clear()
            traiter_formatage_inline(p, texte, doc)
            for run in p.runs:
                run.font.name = POLICE_PRINCIPALE
                run.font.size = Pt(10.5)

        # Séparateur horizontal
        elif ligne.strip() in ("---", "***", "___"):
            ajouter_separateur(doc)

        # Ligne vide
        elif ligne.strip() == "":
            if i > 0 and lignes[i - 1].strip() != "":
                doc.add_paragraph()

        # Paragraphe normal
        else:
            p = doc.add_paragraph()
            traiter_formatage_inline(p, ligne, doc)
            for run in p.runs:
                run.font.name = POLICE_PRINCIPALE
                run.font.size = Pt(10.5)

        i += 1


def generer_docx(fichier_md: Path, dossier_sortie: Path) -> Path:
    """Génère le fichier .docx depuis le markdown."""
    contenu = fichier_md.read_text(encoding="utf-8")
    meta, corps = parser_frontmatter(contenu)

    doc = Document()

    # Marges de page
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    configurer_styles(doc)

    titre_doc = meta.get("titre", fichier_md.stem)
    ajouter_entete(doc, titre_doc)
    ajouter_pied_de_page(doc)

    convertir_markdown_en_docx(corps, doc, meta)

    # Nom du fichier de sortie
    reference = meta.get("reference", fichier_md.stem)
    nom_fichier = re.sub(r"[^\w\-_]", "_", reference) + ".docx"
    chemin_sortie = dossier_sortie / nom_fichier

    dossier_sortie.mkdir(parents=True, exist_ok=True)
    doc.save(str(chemin_sortie))
    return chemin_sortie


def convertir_en_pdf(chemin_docx: Path) -> Path:
    """Convertit le .docx en PDF via LibreOffice si disponible."""
    chemin_pdf = chemin_docx.with_suffix(".pdf")
    try:
        result = subprocess.run(
            [
                "soffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(chemin_docx.parent),
                str(chemin_docx),
            ],
            capture_output=True,
            text=True,
            timeout=60,
        )
        if result.returncode == 0 and chemin_pdf.exists():
            return chemin_pdf
        else:
            print(f"AVERTISSEMENT: Conversion PDF échouée. Sortie soffice: {result.stderr}")
            return chemin_docx
    except FileNotFoundError:
        print("AVERTISSEMENT: LibreOffice (soffice) non trouvé. Pas de conversion PDF.")
        return chemin_docx
    except subprocess.TimeoutExpired:
        print("AVERTISSEMENT: Délai de conversion PDF dépassé.")
        return chemin_docx


def main():
    parser = argparse.ArgumentParser(
        description="Génère un document AKOMA DIGITAL LTD (.docx) depuis un fichier markdown."
    )
    parser.add_argument("fichier", type=Path, help="Fichier markdown source (.md)")
    parser.add_argument(
        "--sortie",
        type=Path,
        default=None,
        help="Répertoire de sortie (défaut: même répertoire que le fichier source)",
    )
    parser.add_argument(
        "--pdf",
        action="store_true",
        help="Convertir également en PDF via LibreOffice",
    )
    args = parser.parse_args()

    fichier_md = args.fichier.resolve()
    if not fichier_md.exists():
        print(f"ERREUR: Fichier introuvable : {fichier_md}")
        sys.exit(1)

    dossier_sortie = args.sortie.resolve() if args.sortie else fichier_md.parent

    print(f"Génération du document AKOMA depuis : {fichier_md}")
    chemin_docx = generer_docx(fichier_md, dossier_sortie)
    print(f"✅ Document produit : {chemin_docx}")

    if args.pdf:
        print("Conversion en PDF...")
        chemin_final = convertir_en_pdf(chemin_docx)
        if chemin_final.suffix == ".pdf":
            print(f"✅ PDF produit : {chemin_final}")
        else:
            print("⚠️  PDF non produit (LibreOffice requis).")


if __name__ == "__main__":
    main()
